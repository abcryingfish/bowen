from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "del_logic_flow.docx"


FLOW_LANES = [
    (
        "A. 启动与定时器",
        [
            ("init(C) 启动", "写“策略启动测试日志”，进入初始化链路。"),
            ("检查 XML 配置", "create_xml_if_not_exists() 不存在则创建，随后清空预警文件并返回。"),
            ("读取界面参数", "set_param() 写入开始/结束时间、最大买入数、价格类型、预警文件、防重规则等。"),
            ("设置资金账号", "set_account() 读取 account 并执行 C.set_account(g.accID)。"),
            ("启动时清理", "若当天 response 不存在则清空 buy.txt；随后撤掉启动前所有委托。"),
            ("导出成本快照", "export_official_position_costs() 查询持仓和历史成交，写官方/历史成本文件。"),
            ("注册定时器", 'C.run_time("on_timer", "3nSecond", ...) 后循环进入主轮询。'),
        ],
    ),
    (
        "B. 每轮 on_timer",
        [
            ("on_timer(C) 进入", "读取当前时间，先做 15:35 后的空仓成本清理。"),
            ("交易时间过滤", "早于开始时间或晚于结束时间，直接返回，不读预警。"),
            ("处理未完成委托", "check_unfinished_orders() 扫 ORDER，可能撤单、等待 tick、构造补单。"),
            ("读取 buy.txt", "load_tdx_signal() 按 7 列读取：代码、名称、时间、价格、涨跌幅、量、公式。"),
            ("预警为空则返回", "文件不存在、读取异常或 DataFrame 为空时，本轮结束。"),
            ("识别买卖标记", "add_signal_flags()：弱卖标记卖出，五日内六级标记买入。"),
            ("丢弃无关公式", "只保留 is_sell == 1 或 is_buy == 1 的预警行。"),
            ("加载当天 response", "load_tdx_response() 读取 buy_response_YYYYMMDD.txt，不存在则用空表。"),
            ("response 防重", "按配置选择 stock_code + is_sell 或 stock_code + datetime + formula 作为已处理键。"),
            ("没有新增预警则返回", "防重过滤后为空，本轮结束；否则写“新增预警”日志。"),
        ],
    ),
    (
        "C. 逐条预警转委托",
        [
            ("逐行遍历新增预警", "每一行都会先转成 signal_info，后续追加处理结果。"),
            ("最大买入标的数检查", "买入信号且 response 中已成交买入股票数达到上限，则直接生成 skipped。"),
            ("convert_order()", "设置方向：卖出为 24，买入为 23；补市场后缀。"),
            ("涨跌幅过滤", "30/688 超过 19.5%，60/00 超过 9.5%，设置 trade_amount = 0 并跳过。"),
            ("买入分支", "只接受五日内六级；读取总资产和当前持仓市值，目标是补到总资产 2%。"),
            ("买入金额计算", "buy_amount = total_asset * target_ratio - current_value；小于等于 0 则跳过。"),
            ("卖出分支", "查询持仓；没有持仓、可用为 0、成本为空、当前价为空都设置卖出数量为 0。"),
            ("卖出成本与分档", "优先历史成交流水成本，兜底官方成本；150% 卖半，200% 清仓。"),
            ("生成委托日志", "把转换后的 order 写入“生成委托信息”。"),
        ],
    ),
    (
        "D. 盘口执行与结果写回",
        [
            ("do_order(C, order)", "若 trade_amount <= 0，直接返回 skipped。"),
            ("读取追单参数", "默认：撤单等待 1 tick，重新下单等待 2 tick，盘口最大轮数 10，最大滑点比例 0.03。"),
            ("买入追单", "取卖一价量，滑点超限则等待；按金额折算一手整数后提交限价单。"),
            ("卖出追单", "取买一价量，滑点超限则等待；按剩余卖出量提交限价单。"),
            ("提交、等待、撤单", "submit_limit_order() 下单，等待 tick，查询成交量，再撤残单。"),
            ("维护成交状态", "买入成交后更新 position_cost.json；150% 卖半成交后标记 half_sell_done。"),
            ("生成处理结果", "make_process_result() 输出 filled、partial_filled、submitted、skipped 或 failed。"),
            ("追加到 response 表", "add_process_result_to_signal() 加上状态、原因、委托号、成交量、剩余量、处理时间。"),
            ("保存 response", "save_tdx_signal_response() 用 GBK 写回制表符文本。"),
            ("收尾日志", "保存后记录账户、持仓、委托、成交，并再次导出官方/历史成本。"),
        ],
    ),
]


BRANCHES = [
    (
        "买入链路",
        [
            "信号：formula == 五日内六级 才允许买入。",
            "仓位：读取账户总资产，单票目标市值为 总资产 * 单次买入仓位比例，默认 2%。",
            "金额：当前持仓市值已达到目标则跳过；否则按差额生成买入金额。",
            "执行：盘口追单吃卖一，按卖一价和卖一量折算一手整数，成交后维护 position_cost.json。",
        ],
    ),
    (
        "卖出链路",
        [
            "持仓：先确认股票在持仓池且可用数量大于 0。",
            "成本：优先使用历史成交流水成本，失败时用官方持仓成本兜底。",
            "分档：当前价达到成本 150% 时卖出 50%，达到 200% 时清仓。",
            "执行：盘口追单吃买一，按买一量和剩余卖出量循环提交限价单。",
        ],
    ),
    (
        "未完成委托处理",
        [
            "扫描：check_unfinished_orders() 查询 ORDER，筛选未完成委托。",
            "买入：已有部分成交则回填自维护成本；若目标仓位已达成，撤同股票未完成买单。",
            "补单：撤原委托，等待 tick，再按剩余数量构造重试委托。",
            "防重：已处理的委托号进入 g.retried_order_ids，避免重复补单。",
        ],
    ),
    (
        "文件和日志",
        [
            "输入：buy.txt 为通达信预警输入，列包含代码、名称、时间、价格、涨跌幅、量、公式。",
            "输出：buy_response_YYYYMMDD.txt 记录每条信号处理状态。",
            "日志：trade_record_log.txt 记录参数、预警、委托、成交、持仓和异常。",
            "成本：position_cost.json 存买入累计金额、累计股数、成本价和 150% 卖半标记。",
        ],
    ),
]


PARAM_ROWS = [
    ("撤单等待tick数", "1", "do_order() 每次提交限价单后", "下单后等待 1 个盘口 tick，再查询成交量并撤掉剩余未成交委托。"),
    ("重新下单等待tick数", "2", "do_order() 盘口为空、滑点超限、本轮未全部成交后；check_unfinished_orders() 撤原单后", "不立即重试，等待 2 个盘口 tick，让盘口刷新后再取新对手价。"),
    ("盘口最大轮数", "10", "do_order() 买入追单、卖出追单、未完成买入补单循环", "最多尝试 10 轮盘口追单；每轮可能提交一次限价单，也可能因为盘口为空或滑点超限只等待不下单。"),
    ("最大滑点比例", "0.03", "is_slippage_exceeded(signal_price, quote_price, max_slippage_ratio)", "对手价相对预警价偏离超过 3% 时，本轮不下单，等待重新下单 tick 后再看盘口。"),
]


FUNCTION_ROWS = [
    ("init(C)", "初始化配置、账户、response 清理、启动撤单、导出成本，并注册定时器。", "create_xml_if_not_exists, set_param, set_account, cancel_all_orders_on_startup"),
    ("on_timer(C)", "交易时间内读取预警，过滤重复和无效信号，逐条生成委托并写 response。", "check_unfinished_orders, load_tdx_signal, convert_order, do_order"),
    ("convert_order(order_info)", "把预警行转换成交易委托，包含涨跌幅过滤、买入仓位计算、卖出分档计算。", "add_market_suffix, get_total_asset, get_history_position_cost_price"),
    ("do_order(C, order)", "按盘口对手价追单，控制滑点，循环提交限价单、等待成交并撤残单。", "get_best_opposite_quote, submit_limit_order, get_order_traded_volume, cancel_order_if_possible"),
    ("check_unfinished_orders(C)", "处理历史未完成委托，撤单后按剩余数量补单，并维护买入成本。", "is_unfinished_order, build_retry_order, do_order"),
    ("save_tdx_signal_response(df, path)", "标准化 response 列，GBK 写回文本，并导出交易/持仓/成本快照。", "normalize_response_columns, log_all_trade_records, log_current_positions, export_official_position_costs"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))

    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])


def set_table_borders(table, color="D9E0EA"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.bold = True
    return p


def add_small_para(cell_or_doc, text, bold_prefix=None):
    p = cell_or_doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_flow_lane(cell, title, items, start_num):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_shading(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_cell_shading(cell, "F8FAFC")
    title_table = cell.add_table(rows=1, cols=1)
    title_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    title_cell = title_table.cell(0, 0)
    set_cell_shading(title_cell, "314155")
    set_cell_margins(title_cell, 80, 120, 80, 120)
    title_cell.paragraphs[0].add_run(title).bold = True
    title_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for offset, (name, desc) in enumerate(items):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{start_num + offset}. {name}")
        run.bold = True
        run.font.color.rgb = RGBColor(37, 50, 69)
        d = cell.add_paragraph()
        d.paragraph_format.left_indent = Inches(0.18)
        d.paragraph_format.space_after = Pt(5)
        d.add_run(desc)


def style_doc(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 12, 6),
        ("Heading 2", 13, "2E74B5", 10, 5),
        ("Heading 3", 11, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("del.py 交易逻辑流程图")
    r.bold = True
    r.font.size = Pt(22)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.color.rgb = RGBColor(31, 41, 51)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(12)
    sub.add_run("根据 del.py（源码编码：GBK）提取；本文档使用 UTF-8 生成。只描述现有逻辑，不改动业务代码。")


def add_metadata(doc):
    add_heading(doc, "文件定位", 1)
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    set_table_widths(table, [3.55, 3.55])
    values = [
        ("脚本角色", "通达信预警信号到 QMT/xtquant 委托执行的自动交易脚本。"),
        ("主入口", "init(C) 初始化，on_timer(C) 每 3 秒轮询。"),
        ("核心链路", "读预警文件 -> 防重 -> 生成委托 -> 盘口追单 -> 写 response。"),
        ("关键文件", "buy.txt、buy_response_YYYYMMDD.txt、trade_record_log.txt、position_cost.json。"),
    ]
    for cell, (label, value) in zip([c for row in table.rows for c in row.cells], values):
        set_cell_shading(cell, "FBFCFE")
        p = cell.paragraphs[0]
        p.add_run(label + "：").bold = True
        p.add_run(value)


def add_main_flow(doc):
    add_heading(doc, "主流程图", 1)
    doc.add_paragraph("读图顺序：先从 A 泳道完成启动，再进入 B 泳道的每轮轮询；B 过滤出新增预警后，进入 C 逐条转委托，最后由 D 执行并写回 response。")
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    set_table_widths(table, [3.55, 3.55])
    num = 1
    for idx, (title, items) in enumerate(FLOW_LANES):
        cell = table.cell(idx // 2, idx % 2)
        add_flow_lane(cell, title, items, num)
        num += len(items)


def add_branches(doc):
    add_heading(doc, "分支逻辑", 1)
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    set_table_widths(table, [3.55, 3.55])
    for cell, (title, items) in zip([c for row in table.rows for c in row.cells], BRANCHES):
        set_cell_shading(cell, "FBFCFE")
        cell.paragraphs[0].add_run(title).bold = True
        for item in items:
            p = cell.add_paragraph(style=None)
            p.paragraph_format.left_indent = Inches(0.12)
            p.add_run("- " + item)


def add_table_section(doc, title, headers, rows, widths):
    add_heading(doc, title, 1)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(header).bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            cell.paragraphs[0].add_run(str(value))
    set_table_widths(table, widths)


def add_warning(doc):
    add_heading(doc, "注意点", 1)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, "F2BAC1")
    set_table_widths(table, [7.1])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF7F8")
    p = cell.paragraphs[0]
    p.add_run("潜在变量问题：").bold = True
    p.add_run("convert_order() 的 150% 卖出分支使用了 self_cost_info.get('half_sell_done')，但在当前函数片段内没有看到 self_cost_info 的赋值。这里没有修改业务逻辑，只标注为需要人工复核的点。")


def build():
    doc = Document()
    style_doc(doc)
    add_title(doc)
    add_metadata(doc)
    add_main_flow(doc)
    add_branches(doc)
    add_table_section(
        doc,
        "追单参数细节",
        ["参数", "默认值", "代码使用位置", "含义"],
        PARAM_ROWS,
        [1.0, 0.55, 2.0, 2.55],
    )
    add_table_section(
        doc,
        "关键函数索引",
        ["函数", "作用", "主要调用"],
        FUNCTION_ROWS,
        [1.4, 2.45, 2.45],
    )
    add_warning(doc)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
