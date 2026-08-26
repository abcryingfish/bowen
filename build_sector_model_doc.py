from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "outputs" / "sector_peak_valley_ml" / "stage_am_model_document"
OUT_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT_DIR / "板块V2波峰波谷概率模型生成过程说明.docx"

FONT_PATH = Path(r"C:\Windows\Fonts\simhei.ttf")
FONT_BOLD_PATH = Path(r"C:\Windows\Fonts\simhei.ttf")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
GRAY = "666666"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GOLD = "7A5A00"


def font(size: int, bold: bool = False):
    path = FONT_BOLD_PATH if bold else FONT_PATH
    return ImageFont.truetype(str(path), size)


def make_pipeline_image(path: Path) -> None:
    image = Image.new("RGB", (1800, 760), "white")
    draw = ImageDraw.Draw(image)
    boxes = [
        (60, 275, 290, 430, "原始数据\n板块/成分\n复权"),
        (350, 275, 580, 430, "因子构建\n六组因子"),
        (640, 275, 870, 430, "V2目标\n峰谷变化"),
        (930, 275, 1160, 430, "第一层\nLightGBM"),
        (1220, 275, 1450, 430, "第二层\nRidge组合"),
        (1510, 275, 1740, 430, "概率校准\n五类走势"),
    ]
    colors = ["DDEBF7", "E2F0D9", "FFF2CC", "FCE4D6", "E4DFEC", "D9EAD3"]
    for i, (x1, y1, x2, y2, text) in enumerate(boxes):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=22, fill="#" + colors[i], outline="#" + DARK_BLUE, width=4)
        lines = text.split("\n")
        box_font_size = 26 if i == 0 else 30
        line_step = 38 if i == 0 else 52
        y = y1 + (30 if i == 0 else 42)
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=font(box_font_size, True))
            draw.text(((x1 + x2 - (bbox[2] - bbox[0])) / 2, y), line, fill="#" + INK, font=font(box_font_size, True))
            y += line_step
        if i < len(boxes) - 1:
            draw.line((x2 + 12, 352, boxes[i + 1][0] - 12, 352), fill="#" + BLUE, width=8)
            draw.polygon([(boxes[i + 1][0] - 12, 352), (boxes[i + 1][0] - 32, 338), (boxes[i + 1][0] - 32, 366)], fill="#" + BLUE)
    draw.text((60, 85), "板块 V2 波峰/波谷概率模型：从数据到部署的完整链路", fill="#" + INK, font=font(38, True))
    draw.text((60, 150), "训练目标是未来 V2 峰谷强度变化，不是直接预测收益", fill="#" + GRAY, font=font(25))
    image.save(path)


def make_ic_image(path: Path) -> None:
    values = [0.1480, 0.2042, 0.1508, 0.1559, 0.1868, 0.1361]
    labels = ["峰超短", "谷超短", "峰5日", "谷5日", "峰20日", "谷20日"]
    image = Image.new("RGB", (1500, 780), "white")
    draw = ImageDraw.Draw(image)
    draw.text((65, 45), "正式六组组合模型：2023+ 测试期 Rank IC", fill="#" + INK, font=font(34, True))
    x0, y0, width, height = 150, 170, 1200, 470
    draw.line((x0, y0 + height, x0 + width, y0 + height), fill="#555555", width=3)
    draw.line((x0, y0, x0, y0 + height), fill="#555555", width=3)
    maxv = 0.24
    for i in range(5):
        v = i * 0.06
        y = y0 + height - int(v / maxv * height)
        draw.line((x0, y, x0 + width, y), fill="#DDDDDD", width=2)
        draw.text((80, y - 15), f"{v:.2f}", fill="#555555", font=font(20))
    bar_w = 120
    gap = 58
    for i, (label, value) in enumerate(zip(labels, values)):
        x = x0 + 70 + i * (bar_w + gap)
        bh = int(value / maxv * height)
        draw.rectangle((x, y0 + height - bh, x + bar_w, y0 + height), fill="#" + BLUE, outline="#" + DARK_BLUE)
        draw.text((x + 20, y0 + height - bh - 42), f"{value:.3f}", fill="#" + INK, font=font(22, True))
        draw.text((x + 18, y0 + height + 20), label, fill="#" + INK, font=font(22))
    draw.text((150, 700), "Rank IC 评价的是预测分与未来 V2 峰谷变化的横截面排序关系。", fill="#" + GRAY, font=font(22))
    image.save(path)


def make_storage_image(path: Path) -> None:
    image = Image.new("RGB", (1600, 820), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 40), "数据与模型文件分层", fill="#" + INK, font=font(36, True))
    rows = [
        ("原始数据", "index_data_daily / stock_basic_data_daily / stock_adj_daily", "行情与复权"),
        ("成分快照", "sector_information/constituent_snapshots_eligible", "最新成分回填历史"),
        ("因子输入", "sector_peak_valley_ml/factor_groups_v1", "六组板块因子"),
        ("技术子组", "sector_peak_valley_ml/technical_subgroups_v1", "18个技术指标组"),
        ("目标标签", "sector_peak_valley_ml/targets_v1", "V2峰谷变化目标"),
        ("模型文件", "sector_peak_valley_ml/models", "LightGBM/Ridge/概率校准"),
        ("部署输出", "outputs/sector_peak_valley_ml/stage_ar_*_5class", "历史结果与最新快照"),
    ]
    y = 125
    for i, (a, b, c) in enumerate(rows):
        fill = "E8EEF5" if i % 2 == 0 else "F7F9FB"
        draw.rectangle((60, y, 1540, y + 78), fill="#" + fill, outline="#C7D3E0", width=2)
        draw.text((90, y + 20), a, fill="#" + DARK_BLUE, font=font(25, True))
        draw.text((330, y + 20), b, fill="#" + INK, font=font(21))
        draw.text((1170, y + 20), c, fill="#" + GRAY, font=font(21))
        y += 82
    image.save(path)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, bold=False, color="000000", font_name="SimHei"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, size=11, bold=False, color="000000", align=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run(r, size={1: 16, 2: 13, 3: 12}[level], bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_bold_label_paragraph(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(label)
    set_run(r, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run(r)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(str(text))
        set_run(r, size=10, bold=True, color=INK)
        set_cell_shading(hdr[i], LIGHT_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(value))
            set_run(r, size=9.5, color=INK)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_image(doc, path, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    p.add_run().add_picture(str(path), width=Inches(width))


def configure_document(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "SimHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for level, size, color, before, after in [(1, 16, BLUE, 16, 8), (2, 13, BLUE, 12, 6), (3, 12, DARK_BLUE, 8, 4)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "SimHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("板块 V2 波峰/波谷概率模型说明")
    set_run(r, size=9, color=GRAY)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("内部研究文档 | 版本：V2标签 + 模型v1 + 五类概率校准v2")
    set_run(r, size=8.5, color=GRAY)


def build_doc():
    pipeline = OUT_DIR / "pipeline.png"
    ic = OUT_DIR / "test_ic.png"
    storage = OUT_DIR / "storage.png"
    make_pipeline_image(pipeline)
    make_ic_image(ic)
    make_storage_image(storage)

    doc = Document()
    configure_document(doc)

    add_text(doc, "板块 V2 波峰/波谷概率模型", size=25, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_text(doc, "从数据、因子、训练、组合到五类走势概率的完整生成过程", size=14, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
    add_text(doc, "适用范围：板块横截面预测 | 训练期：2016-2022 | 测试期：2023年至今", size=10.5, bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=25)
    add_image(doc, pipeline, 6.35)
    add_text(doc, "核心结论：模型预测的不是简单的‘涨/跌’，而是未来 V2 波峰和波谷强度变化所对应的五类走势概率。", size=12, bold=True, color=INK, after=12)
    add_text(doc, "五类最终状态为：波谷看涨、波峰看跌、双向高波、横盘看涨、横盘看跌。概率输出保留了不确定性，最高概率状态只作为辅助解释。", size=11, color=GRAY, after=10)

    add_heading(doc, "一、模型目标与总体设计", 1)
    add_heading(doc, "1.1 模型真正预测什么", 2)
    add_text(doc, "本模型不把未来收益直接作为主要训练目标，而是使用 V2 标签构造未来波峰/波谷强度变化目标。这样做的原因是，收益还会受到仓位、止盈止损、交易成本和持有规则影响；V2 峰谷变化更接近‘未来形态是否向波峰或波谷发展’这一研究问题。")
    add_table(doc, ["周期", "目标定义", "用途"], [
        ("超短", "0.5×(t+1) + 0.3×(t+2) + 0.2×(t+3) - 当前V2", "捕捉短线拐点变化"),
        ("5日", "V2[t+5] - V2[t]", "短周期确认"),
        ("20日", "V2[t+20] - V2[t]", "中周期趋势确认"),
    ], [1500, 4500, 3360])
    add_heading(doc, "1.2 六组因子到五类概率", 2)
    add_text(doc, "模型采用分层结构，而不是把所有因子直接塞进一个大模型。每个因子组先独立学习，得到组内预测分；然后用第二层 Ridge 组合各组分数；最后使用开发期 OOF 预测训练概率校准器。这样可以保留因子组的解释性，并降低不同尺度因子直接混合带来的不稳定。")

    add_heading(doc, "二、原始数据与数据口径", 1)
    add_image(doc, storage, 6.35)
    add_heading(doc, "2.1 原始行情", 2)
    add_table(doc, ["数据", "路径", "作用"], [
        ("板块/指数日线", "D:\\database\\index_data_daily", "板块自身、大盘和市场状态因子"),
        ("成分股日线", "D:\\database\\stock_basic_data_daily", "成分股广度、龙头扩散"),
        ("复权因子", "D:\\database\\stock_adj_daily\\adj_factor_daily", "成分股后复权价格计算"),
        ("成分股快照", "D:\\database\\sector_information\\constituent_snapshots_eligible", "确定板块成员"),
    ], [1800, 4400, 3160])
    add_heading(doc, "2.2 成分股历史口径", 2)
    add_text(doc, "当前没有完整可用的历史成分股数据，因此使用最新成分股快照回填全部历史。该口径保证了训练和推理结构一致，但会带来幸存者偏差风险。成分股相关因子在结果中必须理解为‘最新成分集合下的历史代理’，不是严格历史成分回测。")
    add_heading(doc, "2.3 数据时间切分", 2)
    add_table(doc, ["阶段", "时间", "是否用于调参/定权"], [
        ("开发期", "2016-01-01 至 2022-12-31", "可以"),
        ("测试期", "2023-01-01 至最新数据", "不可以，仅最终评价"),
    ], [1800, 3600, 3960])

    add_heading(doc, "三、因子组生成", 1)
    add_heading(doc, "3.1 技术面组：18个技术子组，427个有效因子", 2)
    add_text(doc, "技术面组是一个‘子组模型的组合’，不是只使用 MACD、RSI 的小模型。技术指标先按指标家族拆成 18 个子组，每个子组独立计算并审计，再由第一层模型和技术大组 Ridge 合成。经过常量、非法字段和覆盖率检查后，最终使用 427 个有效技术因子。")
    add_table(doc, ["技术子组", "技术含义", "存储"], [
        ("ADX / AROON", "趋势强度和趋势方向", "ADX.parquet / AROON.parquet"),
        ("MACD / APO / PPO", "趋势、动量和差离", "对应子组 Parquet"),
        ("RSI / CMO / WILLR / STOCH", "超买超卖和摆动", "对应子组 Parquet"),
        ("BOLL / CCI / ULTOSC", "通道、位置和综合摆动", "对应子组 Parquet"),
        ("MOM / ROC / MFI", "动量、变化率和资金流", "对应子组 Parquet"),
        ("AMA / DEMA / WMA", "自适应和加权均线", "对应子组 Parquet"),
    ], [2100, 4000, 3260])
    add_text(doc, "完整技术子组文件位于 D:\\database\\sector_peak_valley_ml\\technical_subgroups_v1。每个指标子组是一个覆盖全历史的 Parquet 文件，并不是按月拆分。", size=10, color=GRAY)
    add_heading(doc, "3.2 横盘波动组", 2)
    add_text(doc, "包含年化波动率、20/60日波动、ATR区间、20/60日高低位、20日新高占比和新低占比。它主要识别板块是否处于高波动突破、低波动横盘、阶段高位或阶段低位。")
    add_heading(doc, "3.3 相对强弱组", 2)
    add_text(doc, "同时计算板块相对全市场和相对同族板块的强弱，并加入横截面百分位、排名变化和残差强度。核心思想是区分‘市场整体上涨’与‘该板块真正跑赢市场/同族’。")
    add_heading(doc, "3.4 成分股广度组", 2)
    add_text(doc, "使用成分股上涨占比、下跌占比、5日上涨占比、站上20日均线占比、新高新低占比、RSI超买超卖占比，以及有效成分数和覆盖率。它用来判断板块上涨是否由多数成分股扩散支持。")
    add_heading(doc, "3.5 龙头扩散组", 2)
    add_text(doc, "使用 Top5 收益贡献、成交额集中度、成分股中位数收益、收益离散度、龙头与普通成分股收益差、涨跌扩散差，以及涨跌停近似代理。涨跌停字段不是历史严格涨跌停数据，因此同时保留 limit_proxy_coverage 作为覆盖质量标识。")
    add_heading(doc, "3.6 市场状态条件组", 2)
    add_text(doc, "市场状态如果对所有板块相同，不能直接作为有效横截面因子。因此将板块动量、相对强弱、波动、新高新低、成分广度等板块特征分别与大盘收益、波动率、市场宽度和收益离散度相乘，生成 54 个条件化特征。")

    add_heading(doc, "四、第一层模型：各因子组独立训练", 1)
    add_heading(doc, "4.1 每组输出六个目标分", 2)
    add_text(doc, "每个正式因子组针对六个目标分别训练 LightGBM：波峰超短、波谷超短、波峰5日、波谷5日、波峰20日、波谷20日。技术面组先有 18 个子组模型，再合成技术大组；其余组直接输出组分数。")
    add_table(doc, ["层级", "模型数量", "结果"], [
        ("技术子组", "18 × 6", "每个技术指标家族的六个目标分"),
        ("四个非技术核心组", "4 × 6", "横盘、相对强弱、广度、龙头六目标分"),
        ("市场状态组", "1 × 6", "条件化市场状态六目标分"),
        ("技术大组", "Ridge × 6", "把18个技术子组分数合成技术组分数"),
    ], [2200, 1800, 5360])
    add_heading(doc, "4.2 时间隔离与 purge", 2)
    add_text(doc, "训练和验证按时间顺序进行，不随机打乱。为了避免标签未来窗口和滚动窗口在边界处重叠，按目标周期排除训练边界附近的交易日。超短、5日、20日分别使用 43、45、60 个交易日 purge。")
    add_heading(doc, "4.3 第一层 OOF", 2)
    add_text(doc, "开发期内先生成第一层 OOF 预测。OOF 的含义是：每个验证样本的预测来自没有看到该样本未来标签的模型。OOF 不是最终测试结果，而是第二层 Ridge 和概率校准器的合规训练输入。")

    add_heading(doc, "五、第二层 Ridge：合成因子组", 1)
    add_heading(doc, "5.1 先做每日横截面百分位", 2)
    add_text(doc, "不同因子组的原始预测分量纲不同，因此先在每个交易日对每个组的预测分做板块横截面百分位排名，统一到 0 到 1。这样 Ridge 组合的系数反映组间相对贡献，而不是原始数值尺度差异。")
    add_heading(doc, "5.2 严格嵌套 OOF", 2)
    add_text(doc, "第二层不是一次性用全部 OOF 拟合。2021 年验证时只使用更早的合规 OOF；2022 年验证时只使用更早的合规 OOF；2023 年以后测试时，才使用 2020-2022 的合规历史重新拟合最终 Ridge。这解决了‘二层 Ridge 在自身训练样本上评价’的问题。")
    add_heading(doc, "5.3 目标级组选择", 2)
    add_text(doc, "不是所有目标都强行使用市场状态组。根据开发期 OOF 表现，波峰超短、波峰5日和波谷5日使用五个基础组；波谷超短、波峰20日和波谷20日使用六个组。测试期没有参与这个选择。")
    add_table(doc, ["目标", "实际使用组"], [
        ("波峰超短", "技术 + 横盘波动 + 相对强弱 + 成分广度 + 龙头扩散"),
        ("波谷超短", "上述五组 + 市场状态条件"),
        ("波峰5日", "上述五组"),
        ("波谷5日", "上述五组"),
        ("波峰20日", "上述五组 + 市场状态条件"),
        ("波谷20日", "上述五组 + 市场状态条件"),
    ], [1800, 7560])

    add_heading(doc, "六、最终连续分与五类概率", 1)
    add_heading(doc, "6.1 六个连续预测分", 2)
    add_text(doc, "Ridge 输出六个最终连续预测分，分别对应六个 V2 变化目标。之后对每天的波峰分和波谷分分别做横截面百分位，形成 peak_rank 和 valley_rank。方向分定义为：")
    add_text(doc, "direction_score = valley_rank - peak_rank", size=12, bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_text(doc, "方向分越高，表示波谷机会相对强于波峰风险；方向分越低，表示波峰风险相对更强。")
    add_heading(doc, "6.2 五类走势定义", 2)
    add_table(doc, ["波峰排名 P", "波谷排名 V", "状态"], [
        ("P≤0.5；V>0.5", "波峰低；波谷高", "波谷看涨"),
        ("P>0.5；V≤0.5", "波峰高；波谷低", "波峰看跌"),
        ("P>0.5；V>0.5", "波峰高；波谷高", "双向高波"),
        ("P≤0.5；V≤0.5；V>P", "两者低，波谷相对更高", "横盘看涨"),
        ("P≤0.5；V≤0.5；V≤P", "两者低，波峰相对更高或相等", "横盘看跌"),
    ], [2600, 3660, 3100])
    add_heading(doc, "6.3 概率校准", 2)
    add_text(doc, "不能把五类状态硬切成 0/1 后称为概率。当前做法是用开发期 OOF 预测作为输入，训练多分类 Logistic 概率校准器。每个周期单独校准，同时用三个周期的连续特征再训练一个独立共识校准器。共识的周期权重是超短 50%、5日 30%、20日 20%。")
    add_text(doc, "最终建议读取的字段是：prob_consensus_valley_bullish、prob_consensus_peak_bearish、prob_consensus_two_sided_high_volatility、prob_consensus_sideways_bullish、prob_consensus_sideways_bearish。五个字段之和严格为 1。", size=11, bold=True, color=INK)

    add_heading(doc, "七、测试与审计结果", 1)
    add_image(doc, ic, 6.35)
    add_text(doc, "正式六组组合模型的 2023 年以后测试期 Rank IC 如图所示。Rank IC 只评价预测排序和未来 V2 变化排序的一致性，不直接用收益替代 V2 目标。")
    add_table(doc, ["目标", "测试期 Rank IC"], [
        ("波峰超短", "0.1480"), ("波谷超短", "0.2042"), ("波峰5日", "0.1508"),
        ("波谷5日", "0.1559"), ("波峰20日", "0.1868"), ("波谷20日", "0.1361"),
    ], [5000, 4360])
    add_heading(doc, "7.1 概率审计", 2)
    add_table(doc, ["周期", "Log Loss", "Brier", "最高概率命中率"], [
        ("超短", "1.311", "0.680", "43.62%"),
        ("5日", "1.249", "0.656", "45.19%"),
        ("20日", "1.267", "0.663", "44.83%"),
        ("共识", "1.141", "0.615", "49.15%"),
    ], [2200, 2200, 2200, 2760])
    add_text(doc, "当前概率适合表达相对倾向，不适合把 0.50 左右的概率直接解释为高确定性信号。后续若用于交易，还应单独设计持仓、换手、成本和风险控制规则。", size=10.5, color=GRAY)

    add_heading(doc, "八、文件位置与日常使用", 1)
    add_table(doc, ["层级", "路径", "内容"], [
        ("因子组", "D:\\database\\sector_peak_valley_ml\\factor_groups_v1", "六组因子 Parquet"),
        ("技术子组", "D:\\database\\sector_peak_valley_ml\\technical_subgroups_v1", "18个技术子组全历史 Parquet"),
        ("V2目标", "D:\\database\\sector_peak_valley_ml\\targets_v1\\v2_change_targets.parquet", "六个变化目标"),
        ("最终模型", "D:\\database\\sector_peak_valley_ml\\models\\core_blend_oof_selected_v3", "六个 Ridge 模型"),
        ("概率校准", "D:\\database\\sector_peak_valley_ml\\models\\state_probability_v2_5class", "五类校准器"),
        ("最新部署", "outputs\\sector_peak_valley_ml\\stage_ar_deployment_probabilities_5class", "全历史与最新快照"),
    ], [1800, 5200, 2360])
    add_heading(doc, "8.1 日常预测", 2)
    add_text(doc, "日常不需要重新训练模型。流程是：更新最新行情 → 重新生成当日因子 → 调用已经保存的 LightGBM、Ridge 和概率校准器 → 输出最新板块五类概率。")
    add_heading(doc, "8.2 什么时候重新训练", 2)
    add_text(doc, "只有在增加因子、修改 V2 标签、修改预测周期、市场结构明显变化、长期测试表现下降，或进行年度/半年度再训练时，才建立新的模型版本。旧版本应保留，不直接覆盖。")

    add_heading(doc, "九、当前模型的限制", 1)
    add_text(doc, "第一，成分股使用最新快照回填历史，会产生历史成分偏差。第二，涨跌停字段是近似代理，不等同于严格历史涨跌停数据。第三，热点舆情历史过短，尚未接入正式六组组合。第四，基本面聚合组尚未接入。第五，当前结果是离线测试到 2026-06-15，接入实时数据后还需要做持续监控。")
    add_text(doc, "因此，当前模型可以作为‘板块峰谷形态概率研究模型’使用；如果要直接用于交易，还需要增加实时刷新、数据质量监控、概率漂移监控、组合构建和交易成本回测。", size=11, bold=True, color=GOLD)

    add_heading(doc, "附录：当前正式版本标识", 1)
    add_table(doc, ["组件", "版本/名称"], [
        ("目标标签", "V2 change targets"),
        ("六组组合", "v3_oof_selected"),
        ("概率校准", "v2_five_state_calibrated_multinomial"),
        ("部署输出", "v2_five_state_deployment_snapshot"),
    ], [3000, 6360])
    add_text(doc, "本说明文档根据当前实际生成文件、模型清单和测试结果整理。若因子定义、标签定义或训练切分发生变化，应重新生成文档并升级对应版本。", size=10, color=GRAY)

    doc.core_properties.title = "板块 V2 波峰/波谷概率模型生成过程说明"
    doc.core_properties.subject = "板块机器学习模型数据、因子、训练、组合和概率输出流程"
    doc.core_properties.author = "Codex"
    try:
        doc.save(DOCX_PATH)
        saved_path = DOCX_PATH
    except PermissionError:
        # Word 打开旧版本时会锁定原文件；保留旧文件并生成可交付的新版本副本。
        saved_path = DOCX_PATH.with_name(f"{DOCX_PATH.stem}_五类版本{DOCX_PATH.suffix}")
        doc.save(saved_path)
    print(saved_path)


if __name__ == "__main__":
    build_doc()
